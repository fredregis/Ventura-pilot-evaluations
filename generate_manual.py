#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helper functions ──

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1C, 0x28, 0x55)
    return h

def add_screenshot_placeholder(caption):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F0F0',
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n\n[ Screenshot: {caption} ]\n\n')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.italic = True
    doc.add_paragraph()

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f' - {text}')
    else:
        p.add_run(text)
    return p

def add_tip(text):
    p = doc.add_paragraph()
    run = p.add_run('Tip: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x1C, 0x28, 0x55)
    p.add_run(text)

# ── Title Page ──

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Ventura Pilot Evaluations')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1C, 0x28, 0x55)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('User Manual')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('Ventura Air Services')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1C, 0x28, 0x55)

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('May 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ── Table of Contents ──

add_heading_styled('Table of Contents', level=1)
toc_items = [
    '1. Overview',
    '2. Getting Started',
    '3. Creating a New Evaluation',
    '4. Evaluation Detail Screen',
    '5. Grading Items',
    '6. Flight Log',
    '7. Completing an Evaluation',
    '8. PDF Report',
    '9. Export and Import',
    '10. Settings',
    '11. Tips and Best Practices',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ── 1. Overview ──

add_heading_styled('1. Overview', level=1)
doc.add_paragraph(
    'Ventura Pilot Evaluations is an iOS application developed for Ventura Air Services '
    'to evaluate pilots during their Initial Operating Experience (IOE), PIC Upgrades, and '
    'Random evaluations. The app replaces the paper-based VAS Pilot Evaluation Worksheet, '
    'allowing evaluators to grade pilots across 12 sections covering all phases of flight operations.'
)
doc.add_paragraph(
    'The application supports multi-day, multi-session evaluations, flight logging, '
    'PDF report generation, and cross-device evaluation transfer via export and import.'
)

add_screenshot_placeholder('Splash Screen')

doc.add_page_break()

# ── 2. Getting Started ──

add_heading_styled('2. Getting Started', level=1)
doc.add_paragraph(
    'When the app launches, you will see the Ventura Air Services splash screen followed '
    'by the main Pilot Evaluations screen.'
)

add_heading_styled('Main Screen', level=2)
doc.add_paragraph(
    'The main screen displays all evaluations organized into two sections:'
)
add_bullet('Evaluations that are currently being worked on', bold_prefix='In Progress')
add_bullet('Evaluations that have been finalized', bold_prefix='Completed')

doc.add_paragraph()
doc.add_paragraph('Each evaluation row shows:')
add_bullet('Pilot name and status badge (New, In Progress, Session #, or Complete)')
add_bullet('Evaluation type (IOE, PIC Upgrade, or Random)')
add_bullet('Position (PIC or SIC) and aircraft type')
add_bullet('Date created')
add_bullet('Progress bar (orange if unsatisfactory items exist, green otherwise)')

add_screenshot_placeholder('Main Screen - Pilot Evaluations List')

add_heading_styled('Navigation', level=2)
add_bullet('Opens Settings', bold_prefix='Gear icon (top left)')
add_bullet('Opens a menu to create a new evaluation or import one', bold_prefix='Plus icon (top right)')
add_bullet('Swipe left on any evaluation to delete it')

add_screenshot_placeholder('Main Screen - Navigation Elements')

doc.add_page_break()

# ── 3. Creating a New Evaluation ──

add_heading_styled('3. Creating a New Evaluation', level=1)
doc.add_paragraph('To create a new evaluation:')

p = doc.add_paragraph(style='List Number')
p.add_run('Tap the ').bold = False
run = p.add_run('+')
run.bold = True
p.add_run(' icon and select ')
run = p.add_run('New Evaluation')
run.bold = True

p = doc.add_paragraph('Fill in the form:', style='List Number')
p = doc.add_paragraph('Tap ', style='List Number')
run = p.add_run('Create')
run.bold = True
p.add_run(' to save')

doc.add_paragraph()
doc.add_paragraph('The form includes the following fields:')

table = doc.add_table(rows=7, cols=2)
table.style = 'Light Shading Accent 1'
headers = table.rows[0].cells
headers[0].text = 'Field'
headers[1].text = 'Description'
rows_data = [
    ('Evaluation Type', 'Select IOE, PIC Upgrade, or Random (segmented control)'),
    ('Evaluator', 'Enter the evaluator\'s name'),
    ('Last Name / First Name / M.I.', 'Enter the pilot\'s name (last and first name required)'),
    ('Position', 'Select PIC or SIC (segmented control)'),
    ('Aircraft Type', 'Select Citation or Challenger (segmented control)'),
    ('N Number(s)', 'Optional - Enter the aircraft N number(s)'),
]
for i, (field, desc) in enumerate(rows_data):
    row = table.rows[i + 1].cells
    row[0].text = field
    row[1].text = desc

doc.add_paragraph()
add_screenshot_placeholder('New Evaluation Form')

doc.add_page_break()

# ── 4. Evaluation Detail Screen ──

add_heading_styled('4. Evaluation Detail Screen', level=1)
doc.add_paragraph(
    'Tap on any evaluation to open its detail screen. This screen is the hub for all '
    'evaluation activities.'
)

add_screenshot_placeholder('Evaluation Detail Screen - In Progress')

add_heading_styled('Action Buttons', level=2)
doc.add_paragraph('For in-progress evaluations, the following action buttons are available:')

add_bullet(
    'Opens the grading form to evaluate all 63 items across 12 sections.',
    bold_prefix='Start Grading / Continue Grading'
)
add_bullet(
    'Opens the flight log to record flights associated with this evaluation.',
    bold_prefix='Flight Log'
)
add_bullet(
    'Appears after all items are graded if any items received a grade of 1 or 2. '
    'Starts a new session showing only unsatisfactory items for re-evaluation.',
    bold_prefix='Review Unsatisfactory Items'
)
add_bullet(
    'Appears after all items are graded. Prompts for the certifying pilot\'s name to finalize.',
    bold_prefix='Complete Evaluation'
)

add_heading_styled('Pilot Information', level=2)
doc.add_paragraph(
    'Displays all pilot and evaluation details: evaluation type, evaluator name, pilot name, '
    'position, aircraft type, N numbers, date created, and current session number.'
)

add_heading_styled('Progress', level=2)
doc.add_paragraph('Shows a progress bar and grade breakdown:')
add_bullet('Proficient (green) - Items graded 3')
add_bullet('Unsatisfactory (orange) - Items graded 1 or 2')
add_bullet('N/A (gray) - Items marked Not Applicable')
add_bullet('Not Evaluated (gray) - Items marked Not Evaluated')

add_screenshot_placeholder('Progress Section with Grade Breakdown')

add_heading_styled('Comments', level=2)
doc.add_paragraph(
    'Displays all items that have comments, showing the element code, grade, item name, '
    'and comment text.'
)

doc.add_page_break()

# ── 5. Grading Items ──

add_heading_styled('5. Grading Items', level=1)
doc.add_paragraph(
    'The grading screen displays all 12 evaluation sections with their items. '
    'Sections are expanded by default.'
)

add_screenshot_placeholder('Grading Screen - Sections View')

add_heading_styled('Grading Scale', level=2)

table = doc.add_table(rows=6, cols=3)
table.style = 'Light Shading Accent 1'
headers = table.rows[0].cells
headers[0].text = 'Grade'
headers[1].text = 'Meaning'
headers[2].text = 'Color'
grades = [
    ('1', 'Not Proficient', 'Red'),
    ('2', 'Gaining Proficiency', 'Orange'),
    ('3', 'Proficient', 'Green'),
    ('NA', 'Not Applicable', 'Gray'),
    ('NE', 'Not Evaluated', 'Gray'),
]
for i, (grade, meaning, color) in enumerate(grades):
    row = table.rows[i + 1].cells
    row[0].text = grade
    row[1].text = meaning
    row[2].text = color

doc.add_paragraph()

add_heading_styled('How to Grade', level=2)
doc.add_paragraph('For each item, tap a grade button (1, 2, 3, NA, or NE) to assign a grade.')
add_bullet('Tap the same grade again to remove it')
add_bullet('Items graded 1 or 2 (unsatisfactory) automatically show a comment field')
add_bullet('Comments are required for unsatisfactory items')
add_bullet('Tap Save to save and return, or simply navigate back (changes auto-save)')

add_screenshot_placeholder('Grading an Item with Comment Field')

add_heading_styled('Section Controls', level=2)
add_bullet('Tap any section header to collapse or expand it')
add_bullet('Use the collapse/expand button (top left) to toggle all sections at once')
add_bullet('Each section header shows grading progress (e.g., "5/9") and unsatisfactory count')

add_heading_styled('Evaluation Sections (63 items)', level=2)

sections_data = [
    ('1. Preflight (9 items)', 'Crew Briefing, Flight Planning, Weather/NOTAMs/TFR, Exterior Inspection, Cockpit Preflight, W&B/Performance, Fuel Loading & Procedures, Aircraft Appearance, Customer Care'),
    ('2. Before Takeoff (5 items)', 'Takeoff Briefing, NAV/FMS Setup, Flight Director Setup, Engine Start, Taxi'),
    ('3. Takeoff / Climb (6 items)', 'Runway Alignment, Thrust Setting, Speed Control, Crosswind Control, Rotation, Use of Automation'),
    ('4. Cruise (4 items)', 'Use of Automation, Flight Management, Fuel Awareness, Customer Care'),
    ('5. Descent (6 items)', 'Descent Planning, Descent Management, Use of Automation, Approach Briefing, Approach Setup, Customer Care'),
    ('6. Approach / Landing (10 items)', 'Approach Profile, Stabilized Approach, Use of Automation, Speed Control, Flap/Gear Management, Crosswind Control, Touchdown Point, Brakes/Reverse Thrust, Directional Control, Taxi/Parking'),
    ('7. Shutdown / Deplane (3 items)', 'Engine Shutdown, Passenger Deplane, Customer Care'),
    ('8. Post Flight (3 items)', 'Company Communication, Aircraft Appearance, Exterior Inspection'),
    ('9. Securing (4 items)', 'Aircraft Chocked, Gust/Control Lock, Battery Disconnected, Damage Avoidance'),
    ('10. General (5 items)', 'Judgement, CRM, Uniform/Appearance, Timeliness, Required Documents'),
    ('11. Procedures (5 items)', 'High Altitude Airports, Mountainous Airports, Cold Weather Ops, International Procedures, Thunderstorms/Wx Radar'),
    ('12. Company (5 items)', 'Crew Rules & Standards, Ops Specs, SOP\'s, Safety Reporting, Communications'),
]

table = doc.add_table(rows=len(sections_data) + 1, cols=2)
table.style = 'Light Shading Accent 1'
headers = table.rows[0].cells
headers[0].text = 'Section'
headers[1].text = 'Items'
for i, (section, items) in enumerate(sections_data):
    row = table.rows[i + 1].cells
    row[0].text = section
    row[1].text = items

doc.add_paragraph()

add_heading_styled('Multi-Session Evaluations', level=2)
doc.add_paragraph('Evaluations can span multiple days and sessions:')
add_bullet('Grade items over the course of your evaluation period')
add_bullet('After all items are graded, if any are unsatisfactory, the "Review Unsatisfactory Items" button appears')
add_bullet('Tapping it starts a new session and presents only the unsatisfactory items')
add_bullet('This process can repeat across multiple sessions until all items are satisfactory')
add_bullet('The session number is tracked and displayed in the evaluation detail')

add_screenshot_placeholder('Review Unsatisfactory Items Button')

doc.add_page_break()

# ── 6. Flight Log ──

add_heading_styled('6. Flight Log', level=1)
doc.add_paragraph(
    'The flight log allows you to record flights associated with the evaluation.'
)

add_heading_styled('Adding Flights', level=2)
p = doc.add_paragraph('From the evaluation detail screen, tap ', style='List Number')
run = p.add_run('Flight Log')
run.bold = True

p = doc.add_paragraph('Tap ', style='List Number')
run = p.add_run('Add Flight')
run.bold = True
p.add_run(' to add a new entry')

doc.add_paragraph('For each flight, enter:', style='List Number')

doc.add_paragraph()
doc.add_paragraph('Each flight entry contains:')

table = doc.add_table(rows=5, cols=2)
table.style = 'Light Shading Accent 1'
headers = table.rows[0].cells
headers[0].text = 'Field'
headers[1].text = 'Description'
fields = [
    ('Date', 'Select the flight date using the date picker'),
    ('Departure', 'Enter the departure airport (ICAO code)'),
    ('Arrival', 'Enter the arrival airport (ICAO code)'),
    ('Block Time', 'Enter the block time in decimal hours (e.g., 2.5)'),
]
for i, (field, desc) in enumerate(fields):
    row = table.rows[i + 1].cells
    row[0].text = field
    row[1].text = desc

doc.add_paragraph()
add_bullet('The Total Block Time is displayed at the bottom, summing all entries')
add_bullet('Flights are automatically sorted from newest to oldest date')
add_bullet('Swipe left on any flight entry to delete it')
add_bullet('Tap Save to save and return')

add_screenshot_placeholder('Flight Log Screen with Entries')

doc.add_page_break()

# ── 7. Completing an Evaluation ──

add_heading_styled('7. Completing an Evaluation', level=1)
doc.add_paragraph('Once all items have been graded:')

p = doc.add_paragraph(style='List Number')
p.add_run('Review any unsatisfactory items if needed (see Multi-Session Evaluations)')

p = doc.add_paragraph('Tap ', style='List Number')
run = p.add_run('Complete Evaluation')
run.bold = True

p = doc.add_paragraph('Enter the ', style='List Number')
run = p.add_run('certifying pilot\'s name')
run.bold = True
p.add_run(' when prompted')

p = doc.add_paragraph('Tap ', style='List Number')
run = p.add_run('Complete')
run.bold = True
p.add_run(' to finalize')

doc.add_paragraph()
doc.add_paragraph(
    'Once completed, the evaluation moves to the Completed section on the main screen '
    'and can no longer be edited. You can still generate PDF reports and export the evaluation.'
)

add_screenshot_placeholder('Complete Evaluation - Certifying Pilot Prompt')

add_screenshot_placeholder('Completed Evaluation Detail')

doc.add_page_break()

# ── 8. PDF Report ──

add_heading_styled('8. PDF Report', level=1)
doc.add_paragraph(
    'The PDF report is a 3-page document that mirrors the original VAS Pilot Evaluation Worksheet.'
)

add_heading_styled('Page 1 - Evaluation Grid', level=2)
add_bullet('Ventura Air Services logo and company name')
add_bullet('Pilot information (name, position, evaluation type, aircraft, N numbers, evaluator)')
add_bullet('Grading legend')
add_bullet('3-column grid showing all 12 sections with color-coded item grades')

add_screenshot_placeholder('PDF Report - Page 1 (Evaluation Grid)')

add_heading_styled('Page 2 - Comments', level=2)
add_bullet('Element codes, grades, and comments for all items with comments')
add_bullet('Signature area for pilot and certifying pilot')

add_screenshot_placeholder('PDF Report - Page 2 (Comments)')

add_heading_styled('Page 3 - Flight Log', level=2)
add_bullet('Pilot name, aircraft type, and total block time')
add_bullet('Table of all logged flights (date, departure, arrival, block time)')

add_screenshot_placeholder('PDF Report - Page 3 (Flight Log)')

add_heading_styled('Generating and Sharing', level=2)
p = doc.add_paragraph('Tap ', style='List Number')
run = p.add_run('Generate PDF Report')
run.bold = True
p.add_run(' from the evaluation detail screen')

doc.add_paragraph('Preview the PDF in the app', style='List Number')

p = doc.add_paragraph('Tap the ', style='List Number')
run = p.add_run('share icon')
run.bold = True
p.add_run(' (top right) to share via AirDrop, email, print, save to Files, etc.')

add_screenshot_placeholder('PDF Preview with Share Button')

doc.add_page_break()

# ── 9. Export and Import ──

add_heading_styled('9. Export and Import', level=1)
doc.add_paragraph(
    'Evaluations can be transferred between devices using JSON export and import. '
    'This allows an evaluation started on one device to be continued on another.'
)

add_heading_styled('Exporting an Evaluation', level=2)
doc.add_paragraph('Open the evaluation detail screen', style='List Number')
p = doc.add_paragraph('Tap ', style='List Number')
run = p.add_run('Export Evaluation')
run.bold = True
doc.add_paragraph('Choose a sharing method (AirDrop, Messages, email, save to Files, etc.)', style='List Number')

add_screenshot_placeholder('Export Evaluation - Share Sheet')

add_heading_styled('Importing an Evaluation', level=2)
p = doc.add_paragraph('From the main screen, tap the ', style='List Number')
run = p.add_run('+')
run.bold = True
p.add_run(' icon and select ')
run = p.add_run('Import Evaluation')
run.bold = True

doc.add_paragraph('Browse for and select a previously exported JSON file', style='List Number')

p = doc.add_paragraph('You will be asked ', style='List Number')
run = p.add_run('"Add Evaluator?"')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('When prompted to add an evaluator:')
add_bullet('If the evaluation already has an evaluator, the new name will be appended (e.g., "John Smith, Jane Doe")', bold_prefix='Yes')
add_bullet('The evaluation is imported as-is', bold_prefix='No')

doc.add_paragraph()
doc.add_paragraph(
    'A confirmation message appears when the import is successful. If the imported evaluation '
    'has the same ID as an existing evaluation, a new ID is automatically assigned to avoid conflicts.'
)

add_screenshot_placeholder('Import Evaluation - Add Evaluator Prompt')

doc.add_page_break()

# ── 10. Settings ──

add_heading_styled('10. Settings', level=1)
doc.add_paragraph(
    'Tap the gear icon on the main screen to access Settings.'
)
doc.add_paragraph('The settings page displays:')
add_bullet('The current version of the application', bold_prefix='App Version')
add_bullet('The current build number', bold_prefix='Build Number')

add_screenshot_placeholder('Settings Screen')

doc.add_page_break()

# ── 11. Tips ──

add_heading_styled('11. Tips and Best Practices', level=1)

add_tip(
    'You can close the app and return at any time. Your progress is automatically saved.'
)
doc.add_paragraph()

add_tip(
    'Export an evaluation from one device and import it on another to continue '
    'grading with a different evaluator.'
)
doc.add_paragraph()

add_tip(
    'Use the collapse/expand all button in the grading screen to quickly navigate '
    'to the section you need.'
)
doc.add_paragraph()

add_tip(
    'Tap a grade to assign it, tap the same grade again to remove it. '
    'The full grade history is tracked across sessions.'
)
doc.add_paragraph()

add_tip(
    'Always add detailed comments for items graded 1 or 2. These comments appear '
    'on the PDF report and are important for documenting the evaluation.'
)
doc.add_paragraph()

add_tip(
    'Flight logs are sorted from newest to oldest. The total block time is displayed '
    'both in the app and on page 3 of the PDF report.'
)

# ── Save ──

output_path = '/Users/fredericregis/Swift Projects/Ventura Pilot Evaluations/Ventura Pilot Evaluations - User Manual.docx'
doc.save(output_path)
print(f'Manual saved to: {output_path}')
